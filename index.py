import tkinter as tk
from tkinter import ttk
from tkinter import filedialog, messagebox
from docx import Document
from tkcalendar import DateEntry
from docx.shared import RGBColor
from num2words import num2words
from datetime import datetime, timedelta
import locale
import json
import os
import config

# Definir las variables de reemplazo como globales
reemplazos = {}

# Variable global para almacenar la ruta del archivo cargado
archivo_cargado = None



# Cargar el archivo JSON
def cargar_datos_json(ruta_archivo):
    """Carga el archivo JSON y devuelve los datos."""
    with open('municipios.json', 'r', encoding='utf-8') as archivo_json:
        return json.load(archivo_json)
    

def procesar_datos(datos_json):
    """Procesa los datos JSON y devuelve una lista de departamentos y un diccionario de municipios."""
    # Extraer los departamentos únicos
    departamentos = sorted(set(dato["departamento"] for dato in datos_json))

    # Crear un diccionario con los municipios por departamento
    municipios_por_departamento = {}
    for dato in datos_json:
        depto = dato["departamento"]
        if depto not in municipios_por_departamento:
            municipios_por_departamento[depto] = []
        municipios_por_departamento[depto].append(dato["municipio"])

    return departamentos, municipios_por_departamento

salario_inicial = None

def actualizar_salario(event):
    global salario_inicial
    if salario_inicial is None:
        try:
            salario_inicial = int(salario_trabajador.get())
        except ValueError:
            salario_inicial = 0

    seleccion = jornada_trabajo.get()
    if seleccion == "TIEMPO COMPLETO":
        nuevo_salario = salario_inicial
    elif seleccion == "MEDIO TIEMPO":
        nuevo_salario = salario_inicial // 2
    elif seleccion == "POR HORAS":
        nuevo_salario = salario_inicial // 230
    else:
        nuevo_salario = salario_inicial

    salario_trabajador.delete(0, tk.END)
    salario_trabajador.insert(0, f"{nuevo_salario}")



def solo_letras(char):
    if char.isalpha() or char == "":
        return True
    else:
        messagebox.showerror("Entrada inválida", "Solo se permiten letras.")
        return False
    
def solo_numeros(char):
    # Verifica si el carácter ingresado es un número o si está vacío (para permitir borrar)
    if char.isdigit() or char == "":
        return True
    else:
        # Si no es un número, muestra un mensaje de error
        messagebox.showerror("Entrada inválida", "Solo se permiten números.")
        return False

def cargar_documento():
    global archivo_cargado
    global archivo_label
    archivo_cargado = filedialog.askopenfilename(
        title="Seleccionar documento Word",
        filetypes=[("Documentos Word", "*.docx")]
    )
    if archivo_cargado:
        archivo_label.config(text=f"Documento cargado: {archivo_cargado}")
    else:
        archivo_label.config(text="No se ha cargado ningún documento.")

def cargar_documento_por_defecto():
    global archivo_cargado
    archivo_cargado = os.path.join(os.getcwd(), "CONTRATO DE TRABAJO INDEFINIDO.docx")
    if os.path.exists(archivo_cargado):
        archivo_label.config(text=f"Documento cargado: {archivo_cargado}")
    else:
        archivo_label.config(text="No se encontró el documento por defecto.")


def reemplazar_texto_en_documento(documento, reemplazos):
    for parrafo in documento.paragraphs:
        for clave, valor in reemplazos.items():
            if clave in parrafo.text:
                print(f"Reemplazando {clave} con {valor} en el párrafo: {parrafo.text}")
                parrafo.text = parrafo.text.replace(clave, valor)


    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        for clave, valor in reemplazos.items():
                            if clave in run.text:
                                print(f"Reemplazando {clave} con {valor} en una celda")
                                run.text = run.text.replace(clave, valor)
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Establecer el color del texto a negro
                                
    

                                
    nombre_archivo = f"Contrato de Trabajo {reemplazos['[TERMINO]']} de {reemplazos['[TRABAJADOR]']}.docx"
    documento.save(nombre_archivo)



def reemplazar_salario_en_documento(doc_path, salario):
    # Convierte el salario a palabras
    if jornada_trabajo.get() == "POR HORAS":
    #"TIEMPO COMPLETO", "MEDIO TIEMPO", "POR HORAS"
        salario_palabras = num2words(salario, lang='es').replace('coma', 'mil')
        salario_texto = f"{salario:,} ({salario_palabras.upper()} PESOS M/CTE POR HORA.)"
    else:
        salario_palabras = num2words(salario, lang='es').replace('coma', 'mil')
        salario_texto = f"{salario:,} ({salario_palabras.upper()} PESOS M/CTE MENSUAL.)"
    
    # Diccionario de reemplazos
    reemplazos = {"[SALARIO]": salario_texto}
    
    # Imprimir el diccionario de reemplazos para verificar
    print("Diccionario de reemplazos:", reemplazos)

    return reemplazos
    
def calcular_fecha_fin(fecha_inicio, duracion):
    locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    fecha_inicio_dt = datetime.strptime(fecha_inicio, '%d/%m/%Y')
    fecha_fin_dt = fecha_inicio_dt + timedelta(days=duracion)
    return fecha_fin_dt.strftime('%d de %B del %Y').upper()



def deshabilitar_duracion_contrato(event):
    selected_option = termino_contrato.get()
    if selected_option in ["INDEFINIDO", "POR DURACION DE OBRA O LABOR"]:
        entrada_duracion_contrato.grid_remove()  # Oculta el campo
        # entrada_duracion_contrato.config(state='disabled')  # Alternativa: Deshabilita el campo
    else:
        entrada_duracion_contrato.grid()  # Muestra el campo
        # entrada_duracion_contrato.config(state='normal')  # Alternativa: Habilita el campo





# Función para validar la duración del período de prueba
def validar_duracion_prueba(event=None):
    termino = termino_contrato.get()
    print(f"Validando duración de prueba para el término: {termino}")
    
    duracion_prueba_str = entrada_duracion_prueba.get()
    
    if not duracion_prueba_str:
        print("La entrada de duración del período de prueba está vacía")
        #messagebox.showerror("Error", "Por favor, ingrese la duración del período de prueba")
        return
    
    try:
        duracion_prueba = int(duracion_prueba_str)
        print(f"Duración de prueba: {duracion_prueba}")
    except ValueError:
        print("Error al convertir la duración de prueba a entero")
        messagebox.showerror("Error", "Por favor, ingrese un valor numérico válido para la duración del período de prueba")
        return

    if termino == "A TÉRMINO FIJO":
        duracion_contrato_str = entrada_duracion_contrato.get()
        if not duracion_contrato_str:
            print("La entrada de duración del contrato está vacía")
            messagebox.showerror("Error", "Por favor, ingrese la duración del contrato")
            return
        try:
            duracion_contrato = int(duracion_contrato_str)
            print(f"Duración del contrato: {duracion_contrato}")
        except ValueError:
            print("Error al convertir la duración del contrato a entero")
            messagebox.showerror("Error", "Por favor, ingrese un valor numérico válido para la duración del contrato")
            return
        if duracion_prueba > duracion_contrato / 5:
            messagebox.showerror("Error", "No puede exceder la quinta parte de la duración del contrato")
            entrada_duracion_prueba.delete(0, "end")
    elif termino == "INDEFINIDO":
        if duracion_prueba > 60:
            messagebox.showerror("Error", "No puede exceder los 60 días (2 meses) de período de prueba")
            entrada_duracion_prueba.delete(0, "end")
    elif termino == "POR DURACION DE OBRA O LABOR":
        if duracion_prueba > 60:
            messagebox.showerror("Error", "No puede exceder los 60 días (2 meses) de período de prueba")
            entrada_duracion_prueba.delete(0, "end")
    else:
        print("Término del contrato no reconocido")
        
# Función para manejar la selección del primer Combobox
def manejar_seleccion(event):
    deshabilitar_duracion_contrato(event)
    validar_duracion_prueba(event)
    actualizar_objeto_contrato(event)

# Función para actualizar la visibilidad del segundo Combobox
def actualizar_objeto_contrato(event):
    global objeto_contrato
    global termino_objeto_contrato

    if termino_contrato.get() == "POR DURACION DE OBRA O LABOR":        
        objeto_contrato.grid()
        objeto_contrato.grid()
    else:
        objeto_contrato.grid_remove()
        objeto_contrato.grid_remove()

def reemplazar_texto():
    global archivo_cargado
    global reemplazos
    
    fecha_inicio = fecha_inicio_contrato.get()
    
     # Inicializar la duración solo si el término del contrato es "A TÉRMINO FIJO"
    if termino_contrato.get() == "A TÉRMINO FIJO":
        if entrada_duracion_contrato.winfo_ismapped():
            duracion_texto = entrada_duracion_contrato.get()
            if duracion_texto:
                duracion = int(duracion_texto)
            else:
                duracion = 0  # O cualquier valor predeterminado que desees usar
        else:
            duracion = 0  # O cualquier valor predeterminado que desees usar
        fecha_fin = calcular_fecha_fin(fecha_inicio, duracion)
    else:
        fecha_fin = ""  # O cualquier valor predeterminado que desees usar
    

    # Inicializar la lista de campos faltantes
    campos_faltantes = []

    # Verificar si el campo de salario no está vacío
    salario_text = salario_trabajador.get().replace('.', '')
    if not salario_text:
        campos_faltantes.append("Salario")
    else:
        try:
            salario = int(salario_text)
        except ValueError:
            messagebox.showwarning("Advertencia", "El valor del salario no es válido.")
            return

    if archivo_cargado:
        documento = Document(archivo_cargado)
        fecha = fecha_nacimiento.get_date()

        # Obtener los reemplazos de salario solo si salario está definido
        if 'salario' in locals():
            reemplazos_salario = reemplazar_salario_en_documento(archivo_cargado, salario)
        else:
            reemplazos_salario = {}

        locale.setlocale(locale.LC_TIME, 'es_ES')  # Establecer el locale en español

        # Verificar cada campo
        if not entrada_empleador.get():
            campos_faltantes.append("Empleador")
        if not entrada_nit.get():
            campos_faltantes.append("N.I.T del Empleador")
        if not entrada_representante_legal.get():
            campos_faltantes.append("Representante Legal")
        if not entrada_cc_representante_legal.get():
            campos_faltantes.append("C.C. Representante Legal")
        if not entrada_trabajador.get():
            campos_faltantes.append("Trabajador")
        if not entrada_cc_trabajador.get():
            campos_faltantes.append("C.C. Trabajador")
        if not entrada_ciudad.get():
            campos_faltantes.append("Ciudad")
        if not entrada_departamento.get():
            campos_faltantes.append("Departamento")
        if not estado_civil.get():
            campos_faltantes.append("Estado Civil")
        if not entrada_direccion.get():
            campos_faltantes.append("Dirección")
        if not entrada_telefono.get():
            campos_faltantes.append("Teléfono")
        if not entrada_cargo.get():
            campos_faltantes.append("Cargo")
        if not salario_trabajador.get():
            campos_faltantes.append("Salario Base Trabajador")
        if not entrada_departamento_contrato.get():
            campos_faltantes.append("Departamento Contrato")
        if not entrada_ciudad_contrato.get():
            campos_faltantes.append("Ciudad Contrato")        
        if not jornada_trabajo.get():
            campos_faltantes.append("Jornada")
        if not termino_contrato.get():
            campos_faltantes.append("Término del Contrato")
        if not fecha_inicio_contrato.get_date():
            campos_faltantes.append("Fecha de Inicio del Contrato")        
        if not objeto_contrato.get():
            campos_faltantes.append("Objeto del Contrato")
        if not fecha_nacimiento.get_date():
            campos_faltantes.append("Fecha de Nacimiento")        
        if not fecha_firma_contrato.get_date():
            campos_faltantes.append("Fecha de Firma del Contrato")
        if not entrada_duracion_prueba.get():
            campos_faltantes.append("Duración del Período de Prueba")


        # Si hay campos faltantes, mostrar una alerta y no realizar los reemplazos
        if campos_faltantes:
            mensaje_error = "Los siguientes campos están vacíos:\n" + "\n".join(campos_faltantes)
            tk.messagebox.showerror("Error", mensaje_error)
            return

        reemplazos = {
            "[Empleador]": entrada_empleador.get().upper(),
            "[N.I.T]": entrada_nit.get(),
            "[REPRESENTANTE LEGAL]": entrada_representante_legal.get().upper(),
            "[C.C.]" : entrada_cc_representante_legal.get(),
            "[TRABAJADOR]": entrada_trabajador.get().upper(),
            "[C.CNo]": entrada_cc_trabajador.get(),
            "[CIUDAD]": str(entrada_ciudad.get()).upper(),
            "[DEPARTAMENTO]": str(entrada_departamento.get()).upper(),
            "[DIA]": str(fecha.day),
            "[MES]": fecha.strftime('%B').upper(),
            "[ANO]": str(fecha.year),
            "[ESTADO CIVIL]": estado_civil.get().upper(),           
            "[DIRECCION]": entrada_direccion.get().upper(),  
            "[TELEFONO]": entrada_telefono.get(), 
            "[CARGO]": entrada_cargo.get().upper(),
            "[CD_CONT]": str(entrada_ciudad_contrato.get()).upper(),    
            "[DPTO_CONT]": str(entrada_departamento_contrato.get()).upper(),
            "[JORNADA]": jornada_trabajo.get().upper(),
            "[TERMINO]": termino_contrato.get().upper(),
            "[FECHA_INICIO]": fecha_inicio_contrato.get_date().strftime('%d de %B del %Y').upper(),
            "[FECHA_FIN]": fecha_fin.upper(),
            "[FECHA_FIRMA]": fecha_firma_contrato.get_date().strftime('%d de %B del %Y').upper(),
            "[OBJETO]": objeto_contrato.get().upper(),

        }

        # Combinar los diccionarios de reemplazos
        reemplazos.update(reemplazos_salario)

        # Imprimir el diccionario de reemplazos combinado para verificar
        print("Diccionario de reemplazos (combinado):", reemplazos)

        reemplazar_texto_en_documento(documento, reemplazos)

        # Guardar el documento
        archivo_guardado = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documentos Word", "*.docx")]
        )
        if archivo_guardado:
            documento.save(archivo_guardado)
            messagebox.showinfo("Éxito", "El texto ha sido reemplazado y el documento guardado.")
        else:
            messagebox.showwarning("Advertencia", "No se ha guardado el documento.")
    else:
        messagebox.showwarning("Advertencia", "No se ha cargado ningún documento.")

# Definición de la función actualizar_municipios en el ámbito global
def actualizar_municipios(event):
    departamento_seleccionado = entrada_departamento.get()
    entrada_ciudad["values"] = municipios_por_departamento.get(departamento_seleccionado, [])
    entrada_ciudad.set('')  # Limpiar la selección de municipio al cambiar el departamento
        

def actualizar_municipios_contrato(event):
    departamento_seleccionado = entrada_departamento_contrato.get()
    entrada_ciudad_contrato["values"] = municipios_por_departamento.get(departamento_seleccionado, [])
    entrada_ciudad_contrato.set('')  # Limpiar la selección de municipio al cambiar el departamento
        


def create_scrollable_frame(root):
    # Configuración de la ventana principal
    root.geometry("800x600")  # Tamaño inicial
    root.grid_rowconfigure(0, weight=1)  # Hacer la fila 0 expandible
    root.grid_columnconfigure(0, weight=1)  # Hacer la columna 0 expandible

    # Crear un frame contenedor para el canvas y el scrollbar
    container = tk.Frame(root, bg=config.BG_COLOR)
    container.grid(row=0, column=0, sticky="nsew")  # Ocupa todo el espacio

    # Crear un canvas dentro del frame contenedor
    canvas = tk.Canvas(container, bg=config.BG_COLOR)
    canvas.grid(row=0, column=0, sticky="nsew")

    # Hacer que el canvas y el contenedor sean expandibles
    container.grid_rowconfigure(0, weight=1)
    container.grid_columnconfigure(0, weight=1)

    # Crear una scrollbar vertical y asociarla al canvas
    scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
    scrollbar.grid(row=0, column=1, sticky="ns")
    canvas.configure(yscrollcommand=scrollbar.set)

    # Crear un frame dentro del canvas donde se colocarán los widgets
    scrollable_frame = tk.Frame(canvas, bg=config.BG_COLOR)

    # Crear una ventana dentro del canvas para incluir el frame scrollable
    canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")  # Aquí se asigna la ventana a la variable

    # style = ttk.Style()
    # style.configure("Primary.TButton", background=config.PRIMARY_COLOR, foreground="white")
    # style.configure("Secondary.TButton", background=config.SECONDARY_COLOR, foreground="white")
    # style.configure("TFrame", background=config.PRIMARY_COLOR)  # Configurar el color de fondo del frame

    # Función para ajustar el tamaño del canvas cuando cambie el contenido
    def on_frame_configure(event):
        canvas.configure(scrollregion=canvas.bbox("all"))

    # Función para hacer el canvas responsivo al tamaño de la ventana
    def resize_canvas(event):
        canvas_width = event.width
        canvas.itemconfig(canvas_window, width=canvas_width)

    # Función para manejar el evento de la rueda del mouse
    def on_mouse_wheel(event):
        canvas.yview_scroll(int(-1*(event.delta/120)), "units")

    # Asociar eventos
    scrollable_frame.bind("<Configure>", on_frame_configure)
    canvas.bind("<Configure>", resize_canvas) 
    canvas.bind_all("<MouseWheel>", on_mouse_wheel)  # Para Windows y macOS
    canvas.bind_all("<Button-4>", lambda event: canvas.yview_scroll(-1, "units"))  # Para Linux
    canvas.bind_all("<Button-5>", lambda event: canvas.yview_scroll(1, "units"))   # Para Linux

    return scrollable_frame


# Código principal

def main(): 
    global archivo_label  
    global entrada_departamento, entrada_ciudad, entrada_departamento_contrato, entrada_ciudad_contrato 
    global municipios_por_departamento

    # Crear la ventana principal
    root = tk.Tk()
    root.title(config.APP_NAME)
    root.configure(bg=config.BG_COLOR)  # Configura el color de fondo de la ventana principal

    scrollable_frame = create_scrollable_frame(root)

    # Cargar y procesar los datos JSON Municipios y Departamentos Nacimiento
    datos_json = cargar_datos_json('ruta/al/archivo.json')
    departamentos, municipios_por_departamento = procesar_datos(datos_json)

    # # Llamar a la funcion
    # autompletar_municipios(departamentos, municipios_por_departamento)
    

    # Crear el frame desplazable
    scrollable_frame = create_scrollable_frame(root)

    # Definir las variables
    global entrada_empleador, entrada_nit, entrada_representante_legal, entrada_cc_representante_legal
    global entrada_trabajador, entrada_cc_trabajador, entrada_ciudad, entrada_departamento
    global estado_civil, entrada_direccion, entrada_telefono, entrada_cargo, entrada_ciudad_contrato
    global entrada_departamento_contrato, jornada_trabajo, termino_contrato, fecha_inicio_contrato
    global fecha_firma_contrato, objeto_contrato, fecha_nacimiento, salario_trabajador, entrada_duracion_contrato
    global entrada_duracion_prueba

    entrada_empleador = ttk.Entry(scrollable_frame)
    entrada_nit = ttk.Entry(scrollable_frame)
    entrada_representante_legal = ttk.Entry(scrollable_frame)
    entrada_cc_representante_legal = ttk.Entry(scrollable_frame)
    entrada_trabajador = ttk.Entry(scrollable_frame)
    entrada_cc_trabajador = ttk.Entry(scrollable_frame)
    entrada_ciudad = ttk.Entry(scrollable_frame)
    entrada_departamento = ttk.Entry(scrollable_frame)
    estado_civil = tk.StringVar()
    entrada_direccion = ttk.Entry(scrollable_frame)
    entrada_telefono = ttk.Entry(scrollable_frame)
    entrada_cargo = ttk.Entry(scrollable_frame)
    entrada_ciudad_contrato = ttk.Entry(scrollable_frame)
    entrada_departamento_contrato = ttk.Entry(scrollable_frame)
    jornada_trabajo = tk.StringVar()
    termino_contrato = tk.StringVar()
    fecha_inicio_contrato = DateEntry(scrollable_frame)
    fecha_firma_contrato = DateEntry(scrollable_frame)
    objeto_contrato = tk.StringVar()
    fecha_nacimiento = DateEntry(scrollable_frame)
    salario_trabajador = ttk.Entry(scrollable_frame)
    entrada_duracion_contrato = ttk.Entry(scrollable_frame)
    entrada_duracion_prueba = ttk.Entry(scrollable_frame)

    # Configurar las columnas y filas para que se expandan
    for i in range(10):
        scrollable_frame.grid_columnconfigure(i, weight=1)
    for i in range(40):
        scrollable_frame.grid_rowconfigure(i, weight=1)

    # Registrar la función de validación
    vcmd = (root.register(solo_letras), '%P')
    vcmdnum = (root.register(solo_numeros), '%P')

   

    # Subtítulo Datos del Empleador
    tk.Label(scrollable_frame, text="DATOS DEL EMPLEADOR", font=("Arial", 14)).grid(row=1, column=2, columnspan=4, padx=5, pady=10)
    # Datos del Empleador
    tk.Label(scrollable_frame, text="NOMBRE DEL EMPLEADOR", bg=config.BG_LABEL, font=config.FONT_LABEL).grid(row=2, column=1, padx=5, pady=5, sticky="e")
    entrada_empleador = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=config.FONT_ENTRY, validate="key")
    entrada_empleador.grid(row=2, column=2, padx=5, pady=5, sticky="ew")

    tk.Label(scrollable_frame, text="N.I.T EMPLEADOR:", bg=config.BG_LABEL, font=config.FONT_LABEL).grid(row=2, column=3, padx=5, pady=5, sticky="e")
    entrada_nit = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=config.FONT_ENTRY)
    entrada_nit.grid(row=2, column=4, padx=5, pady=5, sticky="ew")

    # Espaciado entre filas
    root.grid_rowconfigure(3, minsize=20)

    # Subtítulo Datos del Representante Legal
    tk.Label(scrollable_frame, text="DATOS DEL REPRESENTANTE LEGAL", font=config.FONT_SUBTITLE).grid(row=4, column=2, columnspan=4, padx=5, pady=10)

    # Datos del Representante Legal
    tk.Label(scrollable_frame, text="REPRESENTANTE LEGAL:", bg=config.BG_LABEL, font=config.FONT_LABEL).grid(row=5, column=1, padx=5, pady=5, sticky="e")
    entrada_representante_legal = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=config.FONT_ENTRY, validate="key")
    entrada_representante_legal.grid(row=5, column=2, padx=5, pady=5, sticky="ew")

    tk.Label(scrollable_frame, text="CC REPRESENTANTE LEGAL:", bg=config.BG_LABEL, font=config.FONT_LABEL).grid(row=5, column=3, padx=5, pady=5, sticky="e")
    entrada_cc_representante_legal = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=config.FONT_ENTRY, validate="key")
    entrada_cc_representante_legal.grid(row=5, column=4, padx=5, pady=5, sticky="ew")

    # Ejecutar la aplicación
    #root.mainloop()

    # Subtítulo Datos del Empleador
    tk.Label(scrollable_frame, text="DATOS DEL EMPLEADOR", font=("Helvetica", 14, "bold")).grid(row=1, column=2, columnspan=4, padx=5, pady=10)

    # Datos del Empleador
    tk.Label(scrollable_frame, text="NOMBRE DEL EMPLEADOR", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=2, column=1, padx=5, pady=5, sticky="e")
    entrada_empleador = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14), validate="key")
    entrada_empleador.grid(row=2, column=2, padx=5, pady=5, sticky="ew")

    tk.Label(scrollable_frame, text="N.I.T EMPLEADOR:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=2, column=3, padx=5, pady=5, sticky="e")
    entrada_nit = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14))
    entrada_nit.grid(row=2, column=4, padx=5, pady=5, sticky="ew")

    # Espaciado entre filas
    root.grid_rowconfigure(3, minsize=20)

    # Subtítulo Datos del Representante Legal
    tk.Label(scrollable_frame, text="DATOS DEL REPRESENTANTE LEGAL", font=("Helvetica", 16, "bold")).grid(row=4, column=2, columnspan=4, padx=5, pady=10)

    # Datos del Representante Legal
    tk.Label(scrollable_frame, text="REPRESENTANTE LEGAL:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=5, column=1, padx=5, pady=5, sticky="e")
    entrada_representante_legal = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14), validate="key")
    entrada_representante_legal.grid(row=5, column=2, padx=5, pady=5, sticky="ew")

    tk.Label(scrollable_frame, text="CC REPRESENTANTE LEGAL:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=5, column=3, padx=5, pady=5, sticky="e")
    entrada_cc_representante_legal = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14), validate="key")
    entrada_cc_representante_legal.grid(row=5, column=4, padx=5, pady=5, sticky="ew")

    # Espaciado entre filas
    root.grid_rowconfigure(6, minsize=20)

    # Subtítulo Datos del Trabajador
    tk.Label(scrollable_frame, text="DATOS DEL TRABAJADOR", font=("Helvetica", 16, "bold")).grid(row=7, column=2, columnspan=4, padx=5, pady=10)

    # Datos del Trabajador
    tk.Label(scrollable_frame, text="NOMBRE TRABAJADOR:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=8, column=1, padx=5, pady=5, sticky="e")
    entrada_trabajador = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14))
    entrada_trabajador.grid(row=8, column=2, padx=5, pady=5, sticky="ew")

    tk.Label(scrollable_frame, text="CC DEL TRABAJADOR:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=8, column=3, padx=5, pady=5, sticky="e")
    entrada_cc_trabajador = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14))
    entrada_cc_trabajador.grid(row=8, column=4, padx=5, pady=5, sticky="ew")

    # Espaciado entre filas
    root.grid_rowconfigure(9, minsize=20)

    # Fecha y lugar de Nacimiento
    tk.Label(scrollable_frame, text="Fecha y lugar de Nacimiento", bg=config.BG_LABEL, font=("Helvetica", 12, "bold")).grid(row=10, column=2, columnspan=4, padx=5, pady=10)

    tk.Label(scrollable_frame, text="Fecha de Nacimiento:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=11, column=1, padx=5, pady=5, sticky="e")
    fecha_nacimiento = DateEntry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14), date_pattern='dd/MM/yyyy')
    fecha_nacimiento.delete(0, "end")
    fecha_nacimiento.insert(0, "dd/MM/AAAA")
    fecha_nacimiento.grid(row=11, column=2, padx=5, pady=5, sticky="ew")

    # Label y combobox para el departamento
    tk.Label(scrollable_frame, text="DEPARTAMENTO:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=11, column=3, padx=5, pady=5, sticky="e")
    entrada_departamento = ttk.Combobox(scrollable_frame, values=departamentos, font=("Helvetica", 14))
    entrada_departamento.grid(row=11, column=4, padx=5, pady=5, sticky="ew")
    entrada_departamento.bind("<<ComboboxSelected>>", actualizar_municipios)

    # Label y combobox para el municipio
    tk.Label(scrollable_frame, text="MUNICIPIO:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=12, column=1, padx=5, pady=5, sticky="e")
    entrada_ciudad = ttk.Combobox(scrollable_frame, font=("Helvetica", 14))
    entrada_ciudad.grid(row=12, column=2, padx=5, pady=5, sticky="ew")
    
    

    

    # Configurar la fuente para los elementos del Combobox
    root.option_add('*TCombobox*Listbox.font', ("Helvetica", 14))
    root.option_add('*TCombobox.font', ("Helvetica", 14)) 

    tk.Label(scrollable_frame, text="ESTADO CIVIL:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=12, column=3, padx=5, pady=5, sticky="e")
    estado_civil = ttk.Combobox(scrollable_frame, values=["SOLTERO", "SOLTERA", "CASADO", "CASADA", "VIUDO", "VIUDA", "SEPARADO", "SEPARADA", "UNION LIBRE"], state="readonly")
    estado_civil.set("Seleccione una opción ...")  # Valor por defecto
    estado_civil.grid(row=12, column=4, padx=5, pady=5, sticky="ew")

    # Dirección
    tk.Label(scrollable_frame, text="Dirección y Teléfono", bg=config.BG_LABEL, font=("Helvetica", 12, "bold")).grid(row=13, column=2, columnspan=4, padx=5, pady=10)

    tk.Label(scrollable_frame, text="DIRECCIÓN:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=14, column=1, padx=5, pady=5, sticky="e")
    entrada_direccion = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14))
    entrada_direccion.grid(row=14, column=2, padx=5, pady=5, sticky="ew")

    tk.Label(scrollable_frame, text="TELÉFONO:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=14, column=3, padx=5, pady=5, sticky="e")
    entrada_telefono = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14))
    entrada_telefono.grid(row=14, column=4, padx=5, pady=5, sticky="ew")

    tk.Label(scrollable_frame, text="TELÉFONO CONTACTO ADICIONAL:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=15, column=1, padx=5, pady=5, sticky="e")
    entrada_telefono = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14))
    entrada_telefono.grid(row=15, column=2, padx=5, pady=5, sticky="ew")

    root.grid_rowconfigure(16, minsize=20)

    # Datos del Contrato
    tk.Label(scrollable_frame, text="DATOS DEL CONTRATO", font=("Helvetica", 16, "bold")).grid(row=17, column=2, columnspan=4, padx=5, pady=10)

    tk.Label(scrollable_frame, text="CARGO QUE DESEMPEÑARÁ:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=18, column=1, padx=5, pady=5, sticky="e")
    entrada_cargo = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14))
    entrada_cargo.grid(row=18, column=2, padx=5, pady=5, sticky="ew")

    tk.Label(scrollable_frame, text="SALARIO BASE DEL TRABAJADOR:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=18, column=3, padx=5, pady=5, sticky="e")
    salario_trabajador = ttk.Entry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14))
    salario_trabajador.grid(row=18, column=4, padx=5, pady=5, sticky="ew")

  
    tk.Label(scrollable_frame, text="DEPARTAMENTO DE LABOR:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=19, column=1, padx=5, pady=5, sticky="e")
    entrada_departamento_contrato = ttk.Combobox(scrollable_frame, values=departamentos, font=("Helvetica", 14))
    entrada_departamento_contrato.grid(row=19, column=2, padx=5, pady=5, sticky="ew")
    entrada_departamento_contrato.bind("<<ComboboxSelected>>", actualizar_municipios_contrato)

    # Label y combobox para el municipio
    tk.Label(scrollable_frame, text="MUNICIPIO DE LABOR:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=19, column=3, padx=5, pady=5, sticky="e")
    entrada_ciudad_contrato = ttk.Combobox(scrollable_frame, font=("Helvetica", 14))
    entrada_ciudad_contrato.grid(row=19, column=4, padx=5, pady=5, sticky="ew")

    # Espaciado entre filas
    root.grid_rowconfigure(18, minsize=20)    


    root.grid_rowconfigure(20, minsize=20)

    tk.Label(scrollable_frame, text="JORNADA DE TRABAJO:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=21, column=1, padx=5, pady=5, sticky="e")
    jornada_trabajo = ttk.Combobox(scrollable_frame, values=["TIEMPO COMPLETO", "MEDIO TIEMPO", "POR HORAS"], state="readonly")
    jornada_trabajo.set("Seleccione una opción ...")  # Valor por defecto
    jornada_trabajo.grid(row=21, column=2, padx=5, pady=5, sticky="ew")
    jornada_trabajo.bind("<<ComboboxSelected>>", actualizar_salario)


    tk.Label(scrollable_frame, text="TÉRMINO DEL CONTRATO:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=21, column=3, padx=5, pady=5, sticky="e")
    termino_contrato = ttk.Combobox(scrollable_frame, values=["INDEFINIDO", "A TÉRMINO FIJO", "POR DURACION DE OBRA O LABOR"], state="readonly")
    termino_contrato.set("Seleccione una opción ...")  # Valor por defecto
    termino_contrato.grid(row=21, column=4, padx=5, pady=5, sticky="ew")
    termino_contrato.bind("<<ComboboxSelected>>", manejar_seleccion)
    termino_contrato.bind("<FocusOut>", validar_duracion_prueba)

    # Espaciado entre filas
    root.grid_rowconfigure(22, minsize=21)

    tk.Label(scrollable_frame, text="OBJETO DEL CONTRATO DE TRABAJO:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=22, column=1, padx=5, pady=5, sticky="e")
    objeto_contrato = ttk.Combobox(scrollable_frame, values=["LICENCIA DE MATERNIDAD", "INCREMENTO DE VENTAS", "VACACIONES"], state="readonly")
    objeto_contrato.set("Seleccione una opción ...")  # Valor por defecto
    objeto_contrato.grid(row=22, column=2, padx=5, pady=5, sticky="ew")



    tk.Label(scrollable_frame, text="Fecha de Inicio de Contrato:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=24, column=1, padx=5, pady=5, sticky="e")
    fecha_inicio_contrato = DateEntry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14), date_pattern='dd/MM/yyyy')
    fecha_inicio_contrato.delete(0, "end")
    fecha_inicio_contrato.insert(0, "dd/MM/AAAA")
    fecha_inicio_contrato.grid(row=24, column=2, padx=5, pady=5, sticky="ew")

    tk.Label(scrollable_frame, text="Fecha de Firma de Contrato:", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=24, column=3, padx=5, pady=5, sticky="e")
    fecha_firma_contrato = DateEntry(scrollable_frame, style="Rounded.TEntry", font=("Helvetica", 14), date_pattern='dd/MM/yyyy')
    fecha_firma_contrato.delete(0, "end")
    fecha_firma_contrato.insert(0, "dd/MM/AAAA")
    fecha_firma_contrato.grid(row=24, column=4, padx=5, pady=5, sticky="ew")

    # Label y combobox para el municipio
    tk.Label(scrollable_frame, text="DURACION DEL CONTRATO (EN DIAS):", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=25, column=1, padx=5, pady=5, sticky="e")
    entrada_duracion_contrato = ttk.Entry(scrollable_frame, font=("Helvetica", 14))
    entrada_duracion_contrato.grid(row=25, column=2, padx=5, pady=5, sticky="ew")
    # Enlazar la función de validación a los eventos de las entradas
    entrada_duracion_contrato.bind("<FocusOut>", validar_duracion_prueba)


    # Label y combobox para el municipio
    tk.Label(scrollable_frame, text="DURACION DEL PERIODO DE PRUEBA (EN DIAS):", bg=config.BG_LABEL, font=("Helvetica", 14, "bold italic")).grid(row=25, column=3, padx=5, pady=5, sticky="e")
    entrada_duracion_prueba = ttk.Entry(scrollable_frame, font=("Helvetica", 14))
    entrada_duracion_prueba.grid(row=25, column=4, padx=5, pady=5, sticky="ew")
    #entrada_duracion_prueba.bind("<FocusOut>", manejar_seleccion)
    entrada_duracion_prueba.bind("<FocusOut>", validar_duracion_prueba)

    # Cargar el documento
    cargar_btn = tk.Button(scrollable_frame, text="Cargar Documento", command=cargar_documento)
    cargar_btn.grid(row=28, column=2, columnspan=2, pady=10, sticky="ew")

    # Botón para reemplazar el texto
    reemplazar_btn = tk.Button(scrollable_frame, text="Reemplazar Texto", command=reemplazar_texto)
    reemplazar_btn.grid(row=29, column=2, columnspan=2, pady=10, sticky="ew")

    # Crear y colocar el Label para mostrar el archivo cargado
    archivo_label = tk.Label(scrollable_frame, text="No se ha cargado ningún documento.")
    archivo_label.grid(row=30, column=2, columnspan=2, pady=10, sticky="ew")

    # Cargar el documento por defecto al iniciar la aplicación
    cargar_documento_por_defecto()

    # # Asegurarse de que el contenedor se expanda
    # root.grid_rowconfigure(0, weight=1)
    # root.grid_columnconfigure(0, weight=1)
    # contenedor.grid_rowconfigure(0, weight=1)
    # contenedor.grid_columnconfigure(0, weight=1)

    # Configuración de la interfaz gráfica y Reinicio de Variables
    #termino_contrato.set("INDEFINIDO")  # Valor por defecto para pruebas


    root.mainloop()
if __name__ == "__main__":
    main()

    